from __future__ import annotations

from transformers import pipeline, AutoTokenizer, AutoModelForTokenClassification
from typing import List, Optional, Any, Tuple, Dict
from enum import Enum
from copy import deepcopy
from docx import Document
from docx.shared import Pt, Cm
from datetime import datetime
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.shared import OxmlElement, qn
from docx.opc import constants
from docx.text.run import Run
from googleapiclient.discovery import build
from google.oauth2 import service_account
from googleapiclient.http import MediaFileUpload
import requests
import re

from run.borrow_detect.borrow import FreqComparator
from pg_prep.pgp_record import GenizaArticle

AR_LABEL = "B-JA"
text = [
    "חצרנא נחן אלשהוד אלכאתמין שהאדתנא אספל הדא אלכתאב פי ביע",
    "עמה ואבאע מנה סתין רטלא בסתה ותלתין דינארא ותבקת",
    "יבק לה עליהם ען הדה אלכמסה ועשרין דינארא שום טענה",
    "נע ותולא קבצה מן אלשיך אלאגל אבי עמר וכילה אלשיך אבי נצר",
    "ור חלפון הלוי החכם והנבון הידוע אבן אלדמיאטי ענד אלשיך אלאגל",
    "וקאל להם אלמשכילים"
]


class Word:
    class Lang(Enum):
        AR = 1,
        NAR = 2,  # Non-Arabic (Hebrew, Aramaic)
        MIX = 3,  # Stem in Hebrew, Article in Arabic
        TBD = 4,  # To Be Decided

    _LABEL_TO_LANG_MAP = {
        "B-JA": Lang.AR,
        "B-NJA": Lang.NAR
    }

    def __init__(self, original_word: str, result_word: str, lang: Lang):
        self._original_word: str = original_word
        self._processed_word: str = result_word
        self._lang: Word.Lang = lang

    def __repr__(self):
        return f"<Word: {self._original_word}, {self._processed_word}, {self._lang.name}>"

    @property
    def original_word(self) -> str:
        return self._original_word

    @property
    def processed_word(self) -> str:
        return self._processed_word

    @property
    def lang(self) -> Lang:
        return self._lang

    @original_word.setter
    def original_word(self, value: str):
        self._original_word = value

    @processed_word.setter
    def processed_word(self, value: str):
        self._processed_word = value

    @lang.setter
    def lang(self, value: Word.Lang):
        self._lang = value

    @staticmethod
    def convert_label(label: str) -> Lang:
        if label not in Word._LABEL_TO_LANG_MAP.keys():
            raise KeyError(f"label {label} is unknown")

        return Word._LABEL_TO_LANG_MAP[label]


class Task:
    _start_time: Optional[datetime]
    _end_time: Optional[datetime]

    def __init__(self):
        self.HE_LETTERS = "אבגדהוזחטיכלמנסעפצקרשתךםןףץ"
        self.MAX_LEN = 510
        self._start_time = datetime.now()
        self._end_time = None

    @staticmethod
    def _is_internal_token(token: str) -> bool:
        return len(token) > 2 and token[:2] == "##"

    def _merge_tokens(self, tokens: List[Dict]):
        raise NotImplementedError

    def get_time_data(self) -> Tuple[datetime, datetime]:
        return self._start_time, self._end_time


class PrePipeline(Task):
    _in: List[str]
    _out: List[List[Word]]

    def __init__(self):
        super().__init__()

    def _process(self) -> List[List[Word]]:
        raise NotImplementedError

    def output(self):
        self._end_time = datetime.now()
        return self._out


class InPipeline(Task):
    TASK_NAME = "token-classification"
    _in: List[List[Word]]
    _out: List[List[Word]]
    _model_name: str

    def __init__(self, inp: List[List[Word]], model_name: Optional[str] = None):
        super().__init__()
        self._in = inp
        self._model_name = model_name

    def _run_nn(self, input_nn: List[str]) -> List[Dict]:
        model = AutoModelForTokenClassification.from_pretrained(self._model_name)
        tokenizer = AutoTokenizer.from_pretrained(self._model_name)
        pipe = pipeline(task=self.TASK_NAME, model=model, tokenizer=tokenizer)

        return pipe(input_nn)

    def output(self):
        return self._out


class PostPipeline(Task):
    _in: List[List[Word]]

    def __init__(self, inp: List[List[Word]]):
        super().__init__()
        self._in = inp

    def _process(self) -> List[List[Word]]:
        raise NotImplementedError


class Import(PrePipeline):
    _out: List[str]

    def __init__(self):
        super().__init__()

    @staticmethod
    def _split_text(text: str) -> List[str]:
        return list(text.splitlines())

    def by_str(self, text: str) -> None:
        if isinstance(text, str) is False:
            raise TypeError("Expected to received str")

        self._out = self._split_text(text)

    def by_list_str(self, text: List[str]) -> None:
        if (isinstance(text, list) and all(isinstance(s, str) for s in text)) is False:
            raise TypeError("Expected to receive [str]")

        self._out = text


    def by_docx_path(self, document_url: str) -> None:
        if isinstance(document_url, str) is False:
            raise TypeError("Expected to received str")
        res = re.compile(r"https:\/\/docs\.google\.com\/document\/d\/.+\/edit\?usp=sharing").match(document_url)
        if res is None:
            raise ValueError("Expected to received URL of Google Doc")

        document_id = document_url.split('/')[-2]
        document_txt_url = f'https://docs.google.com/document/d/{document_id}/export?format=txt'
        response = requests.get(document_txt_url)
        if response.status_code != 200:
            raise RuntimeError(f"The link {document_txt_url} is broken, please check if the permissions are public. Status code: {response.status_code}")

        self._out = self._split_text(response.content.decode("utf-8"))

    def output(self):
        return self._out


class ClearText(PrePipeline):
    _in: List[str]
    _out: List[str]

    def __init__(self, text: List[str]):
        super().__init__()
        self._in = text
        self._out = self._process()

    def _clear_word(self, word) -> str:
        return ''.join(l for l in word if l in self.HE_LETTERS)

    def _clear_text(self, text) -> List[str]:
        return [' '.join(self._clear_word(word) for word in line.strip().split()) for line in text]

    def _process(self) -> List[str]:
        return self._clear_text(self._in)

    def output(self) -> List[str]:
        return self._out


class WrapText(PrePipeline):
    def __init__(self, text: List[str]):
        super().__init__()

        self._in = text
        self._out = self._process()

    def _process(self) -> List[List[Word]]:
        return [[Word(word, "", Word.Lang.TBD) for word in line.split()] for line in self._in]


class CodeSwitch(InPipeline):
    MODEL_NAME = "dwmit/ja_classification"

    def __init__(self, inp: List[List[Word]]):
        super().__init__(inp, model_name=self.MODEL_NAME)
        self._out = self._process()

    def _merge_tokens(self, tokens: Dict) -> List[Word]:
        words: List[Word] = []
        curr_word = ""
        curr_label = ""

        for token in tokens:
            sub_word, lang = token["word"], token["entity"]
            if self._is_internal_token(sub_word):
                curr_word += sub_word[2:]
            else:
                if len(curr_word) > 0:
                    resulted_word = curr_word if Word.convert_label(curr_label) == Word.Lang.NAR else ""
                    words.append(Word(original_word=curr_word, result_word=resulted_word, lang=Word.convert_label(curr_label)))
                    # init
                    curr_word = ""
                curr_word += sub_word
                curr_label = lang

        if len(curr_word) > 0:
            resulted_word = curr_word if Word.convert_label(curr_label) == Word.Lang.NAR else ""
            words.append(Word(original_word=curr_word, result_word=resulted_word, lang=Word.convert_label(curr_label)))

        return words

    def _process(self) -> List[List[Word]]:
        processed_lines = []

        nn_input = [' '.join(word.original_word for word in line) for line in self._in]
        nn_output = self._run_nn(nn_input)

        assert len(nn_output) == len(self._in)
        for i_line in range(len(nn_output)):
            line_output = nn_output[i_line]
            line_result = self._merge_tokens(line_output)
            assert len(line_result) == len(self._in[i_line])
            assert all(
                line_result[i_word].original_word == self._in[i_line][i_word].original_word
                for i_word, word in enumerate(self._in[i_line])
            )
            processed_lines.append(line_result)

        return processed_lines


class BorrowDetector(InPipeline):
    # Detect only words that start with an article (AL) but their stem is in Hebrew
    AR_SUBLINE_PRINT = "ـ"

    PREFIXES = [
        ("ال", "אל"),
        ("لل", "לל"),
        ("لل", "לאל")
    ]

    _freq_comparator: FreqComparator

    def __init__(self, inp: List[List[Word]]):
        super().__init__(inp)
        self._freq_comparator = FreqComparator()
        self._out = self._process()

    def _process(self) -> List[List[Word]]:
        for i_line, line in enumerate(self._in):
            for i_word, word in enumerate(line):
                for prefix_ar, prefix_ja in self.PREFIXES:
                    if word.original_word.startswith(prefix_ja) is False or \
                        len(word.original_word) - len(prefix_ja) <= 2 or \
                        self._freq_comparator.is_mixed(prefix_ar, prefix_ja, word.original_word) is False:
                        continue
                    stem = word.original_word[len(prefix_ja):]
                    word.processed_word = prefix_ar + self.AR_SUBLINE_PRINT + stem
                    word.lang = Word.Lang.MIX

        return self._in


class Transliterate(InPipeline):
    MODEL_NAME = "dwmit/transliterate"

    def __init__(self, inp: List[List[Word]]):
        super().__init__(inp, model_name=self.MODEL_NAME)
        self._out = self._process()

    def _merge_tokens(self, tokens: Dict) -> List[Word]:
        words: List[Word] = []
        curr_word_he, curr_word_ar = "", ""

        for token in tokens:
            letter_he, letter_ar = token["word"], token["entity"][2]
            if self._is_internal_token(letter_he):
                curr_word_he += letter_he[2]
                curr_word_ar += letter_ar
            else:
                assert len(letter_he) == 1
                if len(curr_word_he) > 0:
                    words.append(Word(original_word=curr_word_he, result_word=curr_word_ar, lang=Word.Lang.AR))
                    curr_word_he, curr_word_ar = "", ""
                curr_word_he += letter_he
                curr_word_ar += letter_ar

        if len(curr_word_he) > 0:
            words.append(Word(original_word=curr_word_he, result_word=curr_word_ar, lang=Word.Lang.AR))

        return words

    @staticmethod
    def _merge_ar_he(original_line: List[Word], ar_line: List[Word]):
        merged_line = []

        i_ar = 0
        for original_word in original_line:
            if original_word.lang == Word.Lang.AR:
                merged_line.append(deepcopy(ar_line[i_ar]))
                i_ar += 1
            else:
                merged_line.append(deepcopy(original_word))

        return merged_line

    def _process(self) -> List[List[Word]]:
        processed_lines = []

        nn_input = [' '.join(word.original_word for word in line if word.lang == Word.Lang.AR) for line in self._in]
        nn_output = self._run_nn(nn_input)

        assert len(nn_output) == len(self._in)
        for i_line in range(len(nn_output)):
            line_input = self._in[i_line]
            line_output = nn_output[i_line]
            line_result = self._merge_tokens(line_output)
            line_merged = self._merge_ar_he(line_input, line_result)
            assert len(line_merged) == len(line_input)
            processed_lines.append(line_merged)

        return processed_lines


class SpellingMistakeDetector(InPipeline):
    def __init__(self):
        super().__init__()


class Export(PostPipeline):
    CREDENTIALS_JSON = "../global_def/docx-read-7b56daaf11c4.json"
    LEGAL_OUTPUT_FORMATS = [
        "by_docx_path",
        "by_list_str"
    ]

    _out: str
    _global_start_time: datetime

    def __init__(self, inp: List[List[Word]], *args, **kwargs):
        super().__init__(inp)
        if "global_start_time" not in kwargs:
            raise KeyError("global_start_time hasn't been passed to Export task")
        if "output_format" not in kwargs:
            raise KeyError("output_format hasn't been passed to Export task")

        self._global_start_time = kwargs["global_start_time"]
        self._output_format = kwargs["output_format"]

        if self._output_format == "by_docx_path":
            self._out = self._create_docx()
        elif self._output_format == "by_list_str":
            self._out = self._create_list()
        else:
            raise KeyError(f"output_format {self._output_format} not legal, options: {self.LEGAL_OUTPUT_FORMATS}")

    def _create_list(self):
        return [
            (
                ' '.join([word.original_word for word in geniza_article._processed_words]),
                ' '.join([word.processed_word for word in geniza_article._processed_words])
            )
            for geniza_article in self._in
        ]

    def _add_hyperlink(self, paragraph, text, url):

        # This gets access to the document.xml.rels file and gets a new relation id value
        part = paragraph.part
        r_id = part.relate_to(url, constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        # Create the w:hyperlink tag and add needed values
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        # Create a new run object (a wrapper over a 'w:r' element)
        new_run = Run(OxmlElement('w:r'), paragraph)
        new_run.text = text

        # Join all the xml elements together
        hyperlink.append(new_run._element)
        paragraph._p.append(hyperlink)
        return hyperlink

    def _create_docx(self):

        document = Document()
        h = document.add_heading('Judaeo-Arabic to Arabic transliteration', 0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        table = document.add_table(rows=1, cols=3, style="Table Grid")
        table.autofit = False
        table.allow_autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.columns[0].width = Cm(7)
        table.rows[0].cells[0].width = Cm(7)
        table.columns[1].width = Cm(7)
        table.rows[0].cells[1].width = Cm(7)
        table.columns[2].width = Cm(2)
        table.rows[0].cells[2].width = Cm(2)

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Transliterated'
        hdr_cells[1].text = 'JA'

        hdr_cells[2].text = 'PGPID'  ##'
        for i, geniza_article in enumerate(self._in):

            if len(geniza_article._processed_text) < 5:
                continue
            row_cells = table.add_row().cells

            arabic_cell_par = row_cells[0].paragraphs[0]
            if len(geniza_article._processed_leading_ctxt) > 0:
                ctxt1_run = arabic_cell_par.add_run(geniza_article._processed_leading_ctxt)
            if len(geniza_article._processed_errb) > 0:
                errb_run = arabic_cell_par.add_run(geniza_article._processed_errb)
                errb_run.font.highlight_color = WD_COLOR_INDEX.RED
            if len(geniza_article._processed_errc) > 0:
                errc_run = arabic_cell_par.add_run(geniza_article._processed_errc)
                errc_run.font.highlight_color = WD_COLOR_INDEX.PINK
            if len(geniza_article._processed_target) > 0:
                target_run = arabic_cell_par.add_run(geniza_article._processed_target)
                if len(geniza_article._processed_leading_ctxt) > 0 or len(geniza_article._processed_trailing_ctxt) > 0:
                    target_run.bold = True
                    target_run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
            if len(geniza_article._processed_trailing_ctxt) > 0:
                ctxt2_run = arabic_cell_par.add_run(geniza_article._processed_trailing_ctxt)

            ja_cell_par = row_cells[1].paragraphs[0]
            if len(geniza_article._original_leading_ctxt) > 0:
                ctxt1_run = ja_cell_par.add_run(geniza_article._original_leading_ctxt)
            if len(geniza_article._original_errb) > 0:
                errb_run = ja_cell_par.add_run(geniza_article._original_errb)
                errb_run.font.highlight_color = WD_COLOR_INDEX.RED
            if len(geniza_article._original_target) > 0:
                target_run = ja_cell_par.add_run(geniza_article._original_target)
                if len(geniza_article._original_leading_ctxt) > 0 or len(geniza_article._original_trailing_ctxt) > 0:
                    target_run.bold = True
                    target_run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
            if len(geniza_article._original_trailing_ctxt) > 0:
                ctxt2_run = ja_cell_par.add_run(geniza_article._original_trailing_ctxt)

            pgpid_str = str(geniza_article._pgpid)
            self._add_hyperlink(row_cells[2].paragraphs[0], pgpid_str,
                                f"https://geniza.princeton.edu/en/documents/{pgpid_str}/")

        for i, row in enumerate(table.rows):
            for cell in row.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT if i > 0 else WD_ALIGN_PARAGRAPH.CENTER

        obj_styles = document.styles
        obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.size = Pt(8)
        obj_font.name = 'Times New Roman'

        document.add_paragraph()

        p = document.add_paragraph()
        p.add_run(f"Start time: {self._global_start_time.strftime('%d/%m/%Y %H:%M:%S.%f')[:-3]}", style="CommentsStyle")
        p.add_run().add_break()
        p.add_run(f"End time:  {datetime.now().strftime('%d/%m/%Y %H:%M:%S.%f')[:-3]}", style="CommentsStyle")

        p = document.add_paragraph()
        p.add_run('This tool has been created by ', style="CommentsStyle")
        p.add_run('Daniel Weisberg Mitelman', style="CommentsStyle").bold = True
        p.add_run(' with the supervision of ', style="CommentsStyle")
        p.add_run('Dr. Kfir Bar', style="CommentsStyle").bold = True
        p.add_run(' and ', style="CommentsStyle")
        p.add_run('Prof. Nachum Dershowitz', style="CommentsStyle").bold = True
        p.add_run('.', style="CommentsStyle")

        document.add_page_break()

        final_file_name = f'{datetime.now().strftime("%Y-%m-%d_%H:%M:%S_%f")[:-3]} - JA Transliteration'
        from pathlib import Path
        dir_name = f'../run/transliterations'
        Path(dir_name).mkdir(parents=True, exist_ok=True)
        file_path = f'../run/transliterations/{final_file_name}.docx'

        document.save(file_path)

        credentials = service_account.Credentials.from_service_account_file(
            self.CREDENTIALS_JSON,
            scopes=['https://www.googleapis.com/auth/drive', 'https://www.googleapis.com/auth/documents']
        )

        drive_service = build('drive', 'v3', credentials=credentials)

        file_metadata = {
            'name': 'My Document'
        }
        media = MediaFileUpload(file_path,
                                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        uploaded_file = drive_service.files().create(
            body=file_metadata,
            media_body=media,
            fields='id'
        ).execute()

        drive_file_id = uploaded_file['id']
        conversion_response = drive_service.files().copy(
            fileId=drive_file_id,
            body={
                'parents': [],
                'mimeType': 'application/vnd.google-apps.document',
                'name': final_file_name
            }
        ).execute()

        converted_doc_id = conversion_response['id']

        permission_body = {
            'role': 'writer',
            'type': 'anyone',
            'allowFileDiscovery': False,
        }
        response = drive_service.permissions().create(
            fileId=converted_doc_id,
            body=permission_body,
            fields='id'
        ).execute()

        doc_url = f'https://docs.google.com/document/d/{converted_doc_id}/edit'
        return doc_url

    def output(self):
        return self._out


class PipelineManager:
    _in: List[GenizaArticle]
    _pre_pipeline: List[str]
    _in_pipeline: List[List[Word]]
    _post_pipeline: List[List[Word]]
    _out: str

    PRE_PIPELINE_TASKS = [
        ClearText,
        WrapText
    ]
    IN_PIPELINE_TASKS = [
        CodeSwitch,
        BorrowDetector,
        Transliterate
    ]
    POST_PIPELINE_TASKS = [
        Export
    ]

    def __init__(self, inp: List[str], output_format: str = "by_docx_path", stich_back=True):
        self._in = inp
        self._global_start_time = datetime.now()
        self._output_format = output_format

        self._process(stich_back_long_ones = stich_back)

    def _process_pre_pipeline(self) -> List[List[Word]]:
        for task in self.PRE_PIPELINE_TASKS[:-1]:
            self._pre_pipeline = task(self._pre_pipeline).output()

        return self.PRE_PIPELINE_TASKS[-1](self._pre_pipeline).output()

    def _process_in_pipeline(self) -> List[List[Word]]:
        for task in self.IN_PIPELINE_TASKS:
            self._in_pipeline = task(self._in_pipeline).output()

        return self._in_pipeline

    def _process_post_pipeline(self) -> str:
        for task in self.POST_PIPELINE_TASKS:
            self._out = task(
                self._post_pipeline,
                global_start_time=self._global_start_time,
                output_format=self._output_format
            ).output()
        return self._out

    def _process(self, stich_back_long_ones=True) -> None:

        self._pre_pipeline = [geniza_article._original_text for geniza_article in self._in]
        self._in_pipeline = self._process_pre_pipeline()
        post_pipeline_texts = self._process_in_pipeline()

        # Handling long articles
        prev_article, prev_head, unwanted, idx = None, None, [], 0
        for geniza_article, org_processed_words in zip(self._in, post_pipeline_texts):

            geniza_article.assign_processed(processed_words=org_processed_words)
            #A continuing (long) article?
            if prev_article and prev_article._pgpid == geniza_article._pgpid:

                #Detect and fix duplication and missing pieces
                # geniza_article.detect_and_fix_errors(prev_article)
                # Detect and highlight duplication and missing pieces
                geniza_article.detect_and_highlight_errors(prev_article)

                #Untangle, append and remove
                if stich_back_long_ones:
                    if prev_head and prev_head._pgpid == geniza_article._pgpid:
                        prev_head.merge(geniza_article)
                        unwanted.append(idx)
                    else:
                        prev_head = geniza_article
            else:
                prev_head = geniza_article
            prev_article = geniza_article
            idx = idx + 1

        for i in sorted(unwanted, reverse=True):
            del self._in[i]

        self._post_pipeline = self._in
        self._out = self._process_post_pipeline()

    def output(self):
        return self._out
